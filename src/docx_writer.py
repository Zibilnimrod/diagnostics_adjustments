"""Render the extracted records into a 'טבלת התאמות' .docx.

Layout mirrors the teacher's existing tables: landscape A4, an RTL 4-column
grid, 14pt bold headers.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Emu, Pt

from .extractor import StudentRecord

HEADERS = ["שם התלמיד", "הקשיים", "המלצות להתאמות", "סוג האבחון"]
COLUMN_WIDTHS_DXA = [2060, 4465, 4631, 2953]  # taken from the sample tables

# A4 landscape, matching the samples.
PAGE_WIDTH = Emu(10692130)
PAGE_HEIGHT = Emu(7560310)


def _set_rtl_paragraph(paragraph) -> None:
    """Mark a paragraph and its runs as right-to-left."""
    p_pr = paragraph._p.get_or_add_pPr()
    p_pr.append(OxmlElement("w:bidi"))
    for run in paragraph.runs:
        r_pr = run._r.get_or_add_rPr()
        fonts = OxmlElement("w:rFonts")
        fonts.set(qn("w:hint"), "cs")
        r_pr.append(fonts)
        r_pr.append(OxmlElement("w:rtl"))


def _set_table_rtl(table) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_pr.append(OxmlElement("w:bidiVisual"))


def _write_cell(cell, text: str, bold: bool = False, size: int = 12) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    # Records may carry multi-line values (e.g. several diagnoses).
    for i, line in enumerate((text or "").split("\n")):
        if i > 0:
            paragraph = cell.add_paragraph()
        run = paragraph.add_run(line.strip())
        run.bold = bold
        run.font.size = Pt(size)
        if bold:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_rtl_paragraph(paragraph)


def format_class_title(folder_name: str) -> str:
    """'א2' -> "כיתה א'2". Falls back to the raw folder name."""
    name = folder_name.strip()
    if len(name) >= 2 and name[1:].isdigit():
        return f"כיתה {name[0]}'{name[1:]}"
    return f"כיתה {name}"


def build_class_document(
    records: list[StudentRecord],
    folder_name: str,
    teacher: str | None,
    output_path: Path,
) -> Path:
    document = Document()

    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = PAGE_WIDTH
    section.page_height = PAGE_HEIGHT

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(
        f"תלמידים המקבלים שעות סיוע והתאמות – {format_class_title(folder_name)}"
    )
    run.bold = True
    run.font.size = Pt(16)
    _set_rtl_paragraph(title)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f"מחנכת: {teacher or '____________'}")
    run.font.size = Pt(14)
    _set_rtl_paragraph(subtitle)

    table = document.add_table(rows=1, cols=len(HEADERS))
    table.style = "Table Grid"
    _set_table_rtl(table)

    for cell, header in zip(table.rows[0].cells, HEADERS):
        _write_cell(cell, header, bold=True, size=14)

    for record in records:
        cells = table.add_row().cells
        _write_cell(cells[0], record.student_name)
        _write_cell(cells[1], record.difficulties)
        _write_cell(cells[2], record.accommodations)
        _write_cell(cells[3], record.diagnosis_type)

    for row in table.rows:
        for cell, width in zip(row.cells, COLUMN_WIDTHS_DXA):
            cell.width = Emu(width * 635)  # dxa (twentieths of a point) -> EMU

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))
    return output_path
